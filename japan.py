# импортировать модуль python-docx
from docx import Document 
from docx.shared import Pt

# создать новый документ 
doc1 = Document() 
doc1.add_heading('Япония', level=1) 

# добавить пустой параграф 
par1 = doc1.add_paragraph()

# добавить к параграфу текст с определенными стилевыми настройками 
run1 = par1.add_run('Япония')
run1.font.name = 'Cambria' # установить шрифт Arial
run1.font.size = Pt(14) # установить размер шрифта 24
run1.bold = True # сделать текст полужирным
#run1.italic = True # сделать текст курсивным
#run1.font.underline = True # сделать текст подчеркнутым

run1 = par1.add_run('- это высокогорная страна занимает территорию четрёх крупных и нескольких тысяч мелких ')
run1.font.name = 'Cambria' # установить шрифт Arial
run1.font.size = Pt(14) # установить размер шрифта 24
#run1.bold = True # сделать текст полужирным
#run1.italic = True # сделать текст курсивным
#run1.font.underline = True # сделать текст подчеркнутым

run1 = par1.add_run('островов')
run1.font.name = 'Cambria' # установить шрифт Arial
run1.font.size = Pt(14) # установить размер шрифта 24
#run1.bold = True # сделать текст полужирным
#run1.italic = True # сделать текст курсивным
run1.font.underline = True # сделать текст подчеркнутым

run1 = par1.add_run('. Особенностью природы является преобладание ')
run1.font.name = 'Cambria' # установить шрифт Arial
run1.font.size = Pt(14) # установить размер шрифта 24
#run1.bold = True # сделать текст полужирным
#run1.italic = True # сделать текст курсивным
#run1.font.underline = True # сделать текст подчеркнутым


run1 = par1.add_run('горного рельефа, высокая сейсмичность, активный вулканизм. ')
run1.font.name = 'Cambria' # установить шрифт Arial
run1.font.size = Pt(14) # установить размер шрифта 24
#run1.bold = True # сделать текст полужирным
run1.italic = True # сделать текст курсивным
#run1.font.underline = True # сделать текст подчеркнуты

run1 = par1.add_run('Страна бедна полезными ископаемыми. Из-за вытянутости в меридиональном направлении климатические условаия разноообразны. Более 60% территории, главным образом горы, покрыто лесами: смешанными, широколиственными и преременно-влажными (в том числе ')
run1.font.name = 'Cambria' # установить шрифт Arial
run1.font.size = Pt(14) # установить размер шрифта 24
#run1.bold = True # сделать текст полужирным
#run1.italic = True # сделать текст курсивным
#run1.font.underline = True # сделать текст подчеркнутым

run1 = par1.add_run('муссонными')
run1.font.name = 'Cambria'  # установить шрифт Arial
run1.font.size = Pt(14) # установить размер шрифта 24
#run1.bold = True # сделать текст полужирным
run1.italic = True # сделать текст курсивным
run1.font.underline = True # сделать текст подчеркнуты

run1 = par1.add_run(').')
run1.font.name = 'Cambria'  # установить шрифт Arial
run1.font.size = Pt(14) # установить размер шрифта 24
#run1.bold = True # сделать текст полужирным
#run1.italic = True # сделать текст курсивным
#run1.font.underline = True # сделать текст подчеркнутым

doc1.add_heading('Данные', level=2)

# создать таблицу размером  
table = doc1.add_table(rows=3, cols=2)

# заполнить таблицу данными
data = [ 
        ["Площадь территории", "377944 км2"], 
        ["Плотность населения", "336,3 чел./км2"], 
        ["Население", "126 млн. чел."] 
] # создать вложенный массив с данными

for row_idx, row_data in enumerate(data): 
    row = table.rows[row_idx] # заполнить ряд
    for col_idx, cell_value in enumerate(row_data):
        cell = row.cells[col_idx] 
        cell.text = cell_value 
##
#run2 = par1.add_run(col_idx)
#run2.font.size = Pt(14) # установить размер шрифта 24
#run2.bold = True # сделать текст полужирным
##
# Сохраняем документ 
doc1.save('C:\\Users\\user\\nika_apps\\japan.docx')

